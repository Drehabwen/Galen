from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\DEV\Galen-new")
OUTPUT = ROOT / "docs" / "Galen评测体系测试过程报告.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "64748B"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
GREEN = "1F6B4F"
PALE_GREEN = "EAF5F0"
RED = "9B1C1C"
PALE_RED = "FCECEC"
GOLD = "7A5A00"
PALE_GOLD = "FFF7DB"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, italic=None, color=None, ascii_font="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, color="000000", bold=False):
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths, indent=120):
    assert sum(widths) == 9360, sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页 / 共 ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "NUMPAGES")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rpr.extend([rfonts, color, underline])
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_numbered(doc, title, detail):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(title)
    set_run_font(r, bold=True, color=INK)
    r = p.add_run(f"：{detail}")
    set_run_font(r)
    return p


def add_callout(doc, label, text, kind="info"):
    fill, color = {
        "info": (LIGHT_BLUE, DARK_BLUE),
        "success": (PALE_GREEN, GREEN),
        "risk": (PALE_RED, RED),
        "caution": (PALE_GOLD, GOLD),
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    apply_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    set_run_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_run_font(r, color="222222")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def style_table(table, header=True, font_size=9.5, center_cols=None, status_col=None):
    center_cols = set(center_cols or [])
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if r_idx == 0 and header:
                set_cell_shading(cell, LIGHT_GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in center_cols or r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=font_size, bold=(r_idx == 0), color=INK if r_idx == 0 else "222222")
            if status_col is not None and c_idx == status_col and r_idx > 0:
                value = cell.text.strip()
                set_cell_shading(cell, PALE_GREEN if value == "通过" else PALE_RED)
                for run in cell.paragraphs[0].runs:
                    set_run_font(run, size=font_size, bold=True, color=GREEN if value == "通过" else RED)
    if header:
        set_repeat_table_header(table.rows[0])


def add_source(doc, name, url, note):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    add_hyperlink(p, name, url)
    r = p.add_run(f" - {note}")
    set_run_font(r, size=10)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

# Style sheet: standard_business_brief.
normal = doc.styles["Normal"]
set_style_font(normal, 11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
normal.paragraph_format.widow_control = True

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = doc.styles[name]
    set_style_font(style, size, color, True)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Bullet 2", "List Number"):
    style = doc.styles[list_name]
    set_style_font(style, 11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.167

# Quiet running header and footer.
header = section.header
hp = header.paragraphs[0]
hp.text = "GALEN  |  评测体系测试过程"
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in hp.runs:
    set_run_font(run, size=9, color=MUTED, bold=True)
add_page_field(section.footer.paragraphs[0])

# Memo masthead opening.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("GALEN 评测体系")
set_run_font(r, size=24, bold=True, color=INK)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(18)
r = p.add_run("测试过程报告")
set_run_font(r, size=16, color=DARK_BLUE)

metadata = [
    ("文档状态", "第一阶段 Smoke / 非正式基线"),
    ("测试日期", "2026-08-22"),
    ("被测模型", "DeepSeek V4 Pro"),
    ("代码基点", "galen-research-workbench @ fe4e923；评测实现位于后续本地工作区改动"),
    ("适用范围", "Prompt、上下文工程、Agent Loop、工具调用、Session、Artifact 与内部预览"),
]
for label, value in metadata:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{label}：")
    set_run_font(r, bold=True, color=INK)
    r = p.add_run(value)
    set_run_font(r)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
add_callout(
    doc,
    "核心结论",
    "评测链路已经可运行，但当前数据不足以建立正式基线。4 个真实案例中 3 个通过、1 个因模型请求超预算失败；该失败被原样保留，用于定位 Agent 低效，而不是通过放宽阈值消除。",
    "info",
)

doc.add_page_break()

add_heading(doc, "1. 文档目的与判定边界", 1)
add_body(doc, "本报告记录 Galen 第一阶段评测体系从 GitHub 调研、技术选型、代码实现、案例定义、真实模型运行到结果复核的完整过程。它用于让后续上下文、Prompt、模型参数和 Agent Loop 改动能够相对固定基线进行比较，停止依赖单次运行体感。")
add_callout(doc, "重要边界", "单次 Smoke 只证明链路可执行，不能证明候选版本优于基线。PR Gate 每个 case/model/config 至少运行 5 次；正式 Release 基线建议运行 20-30 次。", "caution")

add_heading(doc, "2. GitHub 调研与技术选型", 1)
add_body(doc, "调研将候选项目分为评测执行器、Agent 轨迹评测、任务环境标准、科研复现基准和观测平台五类。最终选择“借鉴标准、Galen 原生实现”，避免为了评测再引入 Python sidecar、Docker 服务或云端控制面。")

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
headers = ["项目类别", "代表项目", "借鉴内容", "Galen 决策"]
for i, value in enumerate(headers):
    table.rows[0].cells[i].text = value
rows = [
    ("执行器", "Inspect AI / Promptfoo", "Task-Solver-Scorer、声明式 Case、CLI/CI", "借鉴抽象，不嵌入完整框架"),
    ("轨迹", "AgentEvals", "工具轨迹与顺序/集合匹配", "确定性 Rust 断言优先"),
    ("环境", "METR Task Standard", "隔离资产、隐藏评分、环境结束评分", "fixture 复制到临时工作区"),
    ("科研", "CORE-Bench", "产物与可复现结果评分", "加入 Artifact 和事实硬门"),
    ("观测", "Phoenix / Langfuse / Opik", "Trace、实验、数据集、团队标注", "第一阶段不部署；保留 OTEL 扩展位"),
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
apply_table_geometry(table, [1300, 1900, 2860, 3300])
style_table(table, center_cols=[0])

add_heading(doc, "3. 评测体系结构", 1)
add_callout(doc, "执行链路", "CaseSpec  ->  fixture 工作副本  ->  真实 run_chat  ->  RunLedger  ->  硬断言  ->  baseline/candidate 比较", "info")
for title, detail in [
    ("CaseSpec", "使用 TOML 固定输入、预算、必需事实、必需 Artifact、禁止工具和禁止响应模式。"),
    ("隔离 Runner", "只复制 fixture 到临时工作区运行，不直接修改版本化测试原件；失败现场保留用于复盘。"),
    ("RunLedger", "以 JSONL 保存提交、模型、配置哈希、最终响应、完整工具轨迹、Token、分段延迟和工作区位置。"),
    ("硬断言", "先检查事实、路径、工具循环、请求预算和 Artifact，再考虑连续性能指标。"),
    ("比较器", "只比较相同 case/model/config；执行硬门、质量非劣效、P50/P90、Token 和工具错误率判定。"),
]:
    add_numbered(doc, title, detail)

add_heading(doc, "4. 测试环境与准备", 1)
env_table = doc.add_table(rows=1, cols=2)
env_table.style = "Table Grid"
env_table.rows[0].cells[0].text = "项目"
env_table.rows[0].cells[1].text = "配置"
for key, value in [
    ("仓库", r"D:\DEV\Galen-new"),
    ("分支", "galen-research-workbench"),
    ("模型", "deepseek-v4-pro"),
    ("思考等级", "medium；后续机械工具轮关闭深度思考"),
    ("运行入口", "cargo run -p galen --bin eval"),
    ("记录格式", "本地 JSONL；runs 目录默认不提交 Git"),
    ("安全", "报告不包含 API Key；模型凭据仍从本机配置/环境读取"),
]:
    cells = env_table.add_row().cells
    cells[0].text = key
    cells[1].text = value
apply_table_geometry(env_table, [2200, 7160])
style_table(env_table, center_cols=[0])

add_heading(doc, "5. 实施步骤", 1)
steps = [
    ("建立可测性", "扩展 ChatRunSummary，记录上下文组装、MCP、TTFT、TTFR、总耗时、模型请求数、Token 和压缩次数。"),
    ("定义 CaseSpec", "实现 TOML schema、字段默认值、预算校验和 Artifact 路径穿越拒绝。"),
    ("建立隔离工作区", "Runner 在系统临时目录创建工作副本；fixture 中的 plan、evidence 和 GALEN.md 仅作为初始输入。"),
    ("运行真实 Agent Loop", "直接调用 Galen 的 run_chat，而不是只对模型 API 做黑盒问答测试。"),
    ("执行硬断言", "检查运行结束、模型请求预算、工具调用预算、同参数重复、必需工具、必需事实和可预览 Artifact。"),
    ("保存不可变记录", "将完整响应、工具轨迹、断言和工作区位置追加写入 JSONL。"),
    ("比较基线", "只有相同 Case、模型和配置哈希的数据可以比较；不足 5 次返回数据不足。"),
]
for title, detail in steps:
    add_numbered(doc, title, detail)

add_heading(doc, "6. 首批案例", 1)
case_table = doc.add_table(rows=1, cols=4)
case_table.style = "Table Grid"
for i, text in enumerate(["Case", "目标", "主要硬门", "预算"]):
    case_table.rows[0].cells[i].text = text
case_rows = [
    ("E01", "快速科研问答", "包含 FMA-UE；不得写文件或执行命令", "请求 <=2；工具 <=2"),
    ("E04", "多节点计划收敛", "只执行 n3；不重复 n1/n2；生成可预览计划", "请求 <=8；工具 <=12"),
    ("E07", "上下文事实保留", "保留 48、FMA-UE、12 周；生成 context-check.md", "请求 <=6；工具 <=10"),
    ("E09", "Artifact 内部预览", "生成非空 delivery.md，且格式受 Galen 预览支持", "请求 <=6；工具 <=8"),
]
for row in case_rows:
    cells = case_table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
apply_table_geometry(case_table, [900, 1900, 4460, 2100])
style_table(case_table, center_cols=[0, 3])

add_heading(doc, "7. 测试执行过程", 1)
add_heading(doc, "7.1 确定性验证", 2)
add_body(doc, "首先运行案例 schema 校验和 Rust 单元测试，不调用真实模型。这一阶段用于确认 CaseSpec、fixture、轨迹判分和比较器不会因随机输出产生波动。")

cmd_table = doc.add_table(rows=1, cols=2)
cmd_table.style = "Table Grid"
cmd_table.rows[0].cells[0].text = "动作"
cmd_table.rows[0].cells[1].text = "命令 / 结果"
for action, value in [
    ("验证案例", "cargo run -p galen --bin eval -- validate  ->  4 个案例全部通过 schema/fixture 校验"),
    ("运行单元测试", "cargo test -p galen --lib  ->  69 passed, 0 failed"),
    ("敏感信息检查", "新增评测代码、文档和案例未检出明文 sk-... Key"),
    ("差异检查", "git diff --check 通过"),
]:
    cells = cmd_table.add_row().cells
    cells[0].text = action
    cells[1].text = value
apply_table_geometry(cmd_table, [2200, 7160])
style_table(cmd_table, center_cols=[0])

add_heading(doc, "7.2 真实模型 Smoke", 2)
add_body(doc, "随后使用 DeepSeek V4 Pro 分别执行 E01、E04、E07 和 E09。E01 运行两次，用于直接观察同一提交、同一模型、同一案例的自然波动。每次运行都创建独立临时工作区并保存 JSONL。")
add_bullet(doc, "E01-A：首次链路验证；记录 TTFT、TTFR、总耗时和 Token。")
add_bullet(doc, "E01-B：在同一配置下重复，用于观察随机性和网络/推理波动。")
add_bullet(doc, "E04：验证计划节点执行、工具轨迹、收敛和 Artifact。")
add_bullet(doc, "E07：验证工作区记忆中的三个关键事实是否保留。")
add_bullet(doc, "E09：验证交付物是否存在、非空且属于 Galen 支持的预览格式。")

add_heading(doc, "8. 测试结果", 1)
result_table = doc.add_table(rows=1, cols=9)
result_table.style = "Table Grid"
for i, text in enumerate(["Case", "硬门", "TTFT", "TTFR", "总耗时", "请求", "工具/错", "Token", "Artifact"]):
    result_table.rows[0].cells[i].text = text
result_rows = [
    ("E01-A", "通过", "1.034s", "4.165s", "5.811s", "1", "0/0", "3,054", "0/0"),
    ("E01-B", "通过", "1.927s", "9.877s", "11.660s", "1", "0/0", "3,358", "0/0"),
    ("E04", "失败", "0.981s", "1.554s", "27.551s", "9", "8/1", "41,813", "1/1 可预览"),
    ("E07", "通过", "1.285s", "1.702s", "18.514s", "5", "4/0", "23,193", "1/1 可预览"),
    ("E09", "通过", "0.943s", "16.762s", "24.405s", "3", "2/0", "10,777", "1/1 可预览"),
]
for row in result_rows:
    cells = result_table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
apply_table_geometry(result_table, [1050, 700, 800, 800, 850, 700, 850, 1000, 2610])
style_table(result_table, font_size=8.6, center_cols=list(range(9)), status_col=1)

add_callout(doc, "结果摘要", "5 次真实运行中 4 次通过硬门、1 次失败。E04 的失败来自模型请求数 9 次超过预算 8 次，而不是 Artifact 缺失。", "caution")

add_heading(doc, "9. 分案例分析", 1)
add_heading(doc, "9.1 E01 - 响应速度与自然波动", 2)
add_body(doc, "两次 E01 均只发送 1 次模型请求、没有工具调用，并正确包含 FMA-UE。第二次 TTFR 为 9.877 秒，较第一次 4.165 秒高约 137%；总耗时也接近翻倍。")
add_callout(doc, "判定", "不能把单次 Pro 响应速度当作框架性能。至少需要交错运行 5 次形成 PR Gate；正式 P90 结论应基于 20-30 次。", "info")

add_heading(doc, "9.2 E04 - 正确交付但执行低效", 2)
add_body(doc, "E04 正确读取计划并生成 output/research-plan.md，Artifact 非空且可预览；同参数工具调用最大重复次数为 1，没有形成传统意义上的死循环。")
add_bullet(doc, "模型请求 9 次，超过案例上限 8 次。")
add_bullet(doc, "工具调用 8 次，其中 1 次错误。")
add_bullet(doc, "轨迹包含 list_files、read_file、search_files、rehab_data、write_file；rehab_data 对本任务价值较低。")
add_bullet(doc, "总 Token 达 41,813，是本批案例最高值。")
add_callout(doc, "负优化护栏生效", "虽然最终产物正确，评测仍判定失败。不得通过把请求预算从 8 改成 9 来让当前结果通过；下一步应优化工具暴露范围和 Tool Result 后的上下文增量。", "risk")

add_heading(doc, "9.3 E07 - 上下文事实保留", 2)
add_body(doc, "Runner 将包含研究约束的 GALEN.md 复制到独立工作区。最终响应或 Artifact 中保留了 48、FMA-UE 和 12 周三个必需事实，并生成可预览的 output/context-check.md。该结果说明事实断言、工作区记忆读取和 Artifact 验证链路可用。")

add_heading(doc, "9.4 E09 - Artifact 生成与内部预览", 2)
add_body(doc, "E09 生成了非空 output/delivery.md，文件扩展名属于 Galen 预览支持集合，硬门通过。TTFT 为 0.943 秒，但 TTFR 达 16.762 秒，说明模型连接并不慢，首次有效行动较慢。")
add_callout(doc, "诊断价值", "如果只记录 TTFT，会错误得出“响应很快”的结论；TTFR 才能反映用户看到有效正文或工具动作之前的真实等待。", "info")

add_heading(doc, "10. 负优化判定规则", 1)
add_body(doc, "候选版本按以下顺序判定。硬质量门拥有最高优先级，速度和 Token 收益不能抵消事实、状态或交付错误。")
for title, detail in [
    ("硬门", "事实丢失、引用/Artifact 不可追溯、重复执行、工作区串扰、不可预览交付物、权限越界，任一出现即拒绝。"),
    ("质量非劣效", "综合质量相对基线下降超过 3 个百分点即拒绝。"),
    ("有效收益", "TTFR/总耗时/Token 至少改善 15%，或成功率改善至少 5 个百分点，才具备升级基线的理由。"),
    ("尾部保护", "P90 延迟不得恶化超过 10%；工具错误率和重复调用不得增加。"),
    ("数据充分性", "相同 case/model/config 每组不足 5 次时返回数据不足，而不是接受。"),
]:
    add_numbered(doc, title, detail)

doc.add_page_break()

add_heading(doc, "11. 当前结论与限制", 1)
add_callout(doc, "当前结论", "评测基础设施已经能够发现真实低效和速度波动；尚不能宣称当前 Galen 已达到稳定基线。", "success")
add_heading(doc, "11.1 已确认", 2)
for item in [
    "CaseSpec、fixture、真实 run_chat、RunLedger 和硬断言已经贯通。",
    "四个首批案例均可实际执行，而不只是通过格式校验。",
    "E04 的请求预算失败能被自动拦截。",
    "关键事实与可预览 Artifact 能被确定性验证。",
    "TTFT 与 TTFR 分离能够揭示有效行动延迟。",
]:
    add_bullet(doc, item)

add_heading(doc, "11.2 尚未完成", 2)
for item in [
    "每案例 5 次的开发基线与 20-30 次 Release 基线。",
    "关闭应用后的 Session 恢复、跨工作区隔离专用案例。",
    "Evidence 引用真实性和科研方法 Rubric。",
    "A-B-B-A 自动交错调度与 HTML 对比报告。",
    "Galen 内部评测结果预览界面。",
]:
    add_bullet(doc, item)

add_heading(doc, "11.3 工程检查说明", 2)
add_body(doc, "Rust 单元测试 69 项全部通过，案例校验、敏感信息扫描和差异检查通过。严格 Clippy -D warnings 仍会被仓库已有 medical-core、runtime 和既有 Galen 模块警告阻挡；本轮没有把清理全部历史 Clippy 债务扩大到评测实施范围。")

add_heading(doc, "12. 下一阶段执行清单", 1)
for item in [
    "冻结当前四个 CaseSpec 和阈值，阈值修改与产品优化不得位于同一变更。",
    "为 E01、E04、E07、E09 各补足 5 次，并使用交错顺序降低网络时段偏差。",
    "优先分析 E04：缩小工具集合，记录每轮输入 Token 增量，定位无关 rehab_data 调用来源。",
    "建立第一份 baseline JSONL；只有全硬门通过的版本才能成为基线。",
    "任何 Prompt、上下文裁剪或 max-turn 修改均生成 candidate JSONL，并通过 compare 命令判定。",
    "积累 20-30 次后校准自然波动，才正式启用 P90 Release 门槛。",
]:
    add_bullet(doc, f"□ {item}")

add_heading(doc, "附录 A - 常用命令", 1)
commands = [
    ("校验案例", "cargo run -p galen --bin eval -- validate"),
    ("单次 Smoke", "cargo run -p galen --bin eval -- run --case E01 --repeat 1"),
    ("PR Gate", "cargo run -p galen --bin eval -- run --case E01 --repeat 5 --output ../evals/runs/e01-candidate.jsonl"),
    ("比较基线", "cargo run -p galen --bin eval -- compare --baseline <baseline.jsonl> --candidate <candidate.jsonl>"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "用途"
table.rows[0].cells[1].text = "命令"
for label, command in commands:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = command
apply_table_geometry(table, [1900, 7460])
style_table(table, center_cols=[0], font_size=9.2)
for row in table.rows[1:]:
    for run in row.cells[1].paragraphs[0].runs:
        set_run_font(run, size=8.8, ascii_font="Consolas", east_asia="Microsoft YaHei")

add_heading(doc, "附录 B - GitHub 调研来源", 1)
add_source(doc, "Inspect AI", "https://github.com/UKGovernmentBEIS/inspect_ai", "大型语言模型评测框架，提供多轮、工具和模型评分能力。")
add_source(doc, "Promptfoo", "https://github.com/promptfoo/promptfoo", "声明式 LLM/Agent 测试、比较与 CI。")
add_source(doc, "METR Task Standard", "https://github.com/METR/task-standard", "Agent 任务环境、资产、权限和评分标准。")
add_source(doc, "AgentEvals", "https://github.com/langchain-ai/agentevals", "Agent 中间轨迹和工具调用评测。")
add_source(doc, "CORE-Bench", "https://github.com/siegelz/core-bench", "科研论文计算复现 Agent 基准。")
add_source(doc, "Phoenix", "https://github.com/Arize-ai/phoenix", "OpenInference 追踪、实验与评测平台。")
add_source(doc, "Langfuse", "https://github.com/langfuse/langfuse", "LLM 可观测性、数据集、实验和人工标注。")
add_source(doc, "Opik", "https://github.com/comet-ml/opik", "Agent/RAG 追踪、离线评测与监控。")

# Avoid rows splitting across pages where possible.
for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

# Document metadata.
doc.core_properties.title = "Galen评测体系测试过程报告"
doc.core_properties.subject = "Galen Agent evaluation framework implementation and smoke testing"
doc.core_properties.author = "Galen 项目组"
doc.core_properties.keywords = "Galen, Agent Eval, DeepSeek, TTFT, TTFR, Artifact"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
